import os
import PyPDF2
from docx import Document
from typing import Optional, Dict, Any
import aiofiles
from fastapi import UploadFile
import tempfile
import re

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._process_pdf,
            'application/msword': self._process_doc,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'text/plain': self._process_txt
        }

    async def upload_and_process(self, file: UploadFile) -> str:
        """Upload and process a document file"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as temp_file:
                # Write uploaded file to temp file
                content = await file.read()
                temp_file.write(content)
                temp_file.flush()
                
                # Process based on content type
                if file.content_type in self.supported_formats:
                    processor = self.supported_formats[file.content_type]
                    text_content = await processor(temp_file.name)
                else:
                    # Fallback to text processing
                    text_content = content.decode('utf-8', errors='ignore')
                
                # Clean up temp file
                os.unlink(temp_file.name)
                
                return text_content
                
        except Exception as e:
            print(f"Error processing uploaded file: {e}")
            return ""

    async def process_document(self, content: str, document_type: str) -> str:
        """Process document content based on type"""
        try:
            if document_type.lower() == 'pdf':
                return await self._process_pdf_content(content)
            elif document_type.lower() in ['doc', 'docx']:
                return await self._process_word_content(content)
            elif document_type.lower() == 'txt':
                return await self._process_text_content(content)
            else:
                return self._clean_text(content)
                
        except Exception as e:
            print(f"Error processing document: {e}")
            return content

    async def _process_pdf(self, file_path: str) -> str:
        """Process PDF file"""
        try:
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            
            return self._clean_text(text_content)
            
        except Exception as e:
            print(f"Error processing PDF: {e}")
            return ""

    async def _process_doc(self, file_path: str) -> str:
        """Process DOC file (simplified - would need additional libraries for full support)"""
        try:
            # For .doc files, we'll need additional libraries like python-docx2txt
            # For now, return a placeholder
            return "DOC file processing requires additional libraries"
            
        except Exception as e:
            print(f"Error processing DOC: {e}")
            return ""

    async def _process_docx(self, file_path: str) -> str:
        """Process DOCX file"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return self._clean_text(text_content)
            
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return ""

    async def _process_txt(self, file_path: str) -> str:
        """Process TXT file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            return self._clean_text(content)
            
        except Exception as e:
            print(f"Error processing TXT: {e}")
            return ""

    async def _process_pdf_content(self, content: str) -> str:
        """Process PDF content (if already extracted)"""
        return self._clean_text(content)

    async def _process_word_content(self, content: str) -> str:
        """Process Word document content"""
        return self._clean_text(content)

    async def _process_text_content(self, content: str) -> str:
        """Process plain text content"""
        return self._clean_text(content)

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove multiple consecutive line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()

    def extract_metadata(self, content: str) -> Dict[str, Any]:
        """Extract metadata from document content"""
        try:
            metadata = {
                "word_count": len(content.split()),
                "character_count": len(content),
                "paragraph_count": len([p for p in content.split('\n\n') if p.strip()]),
                "sentence_count": len(re.split(r'[.!?]+', content)),
                "average_sentence_length": 0,
                "reading_time_minutes": 0
            }
            
            # Calculate average sentence length
            sentences = re.split(r'[.!?]+', content)
            sentences = [s.strip() for s in sentences if s.strip()]
            if sentences:
                metadata["average_sentence_length"] = sum(len(s.split()) for s in sentences) / len(sentences)
            
            # Estimate reading time (average 200 words per minute)
            metadata["reading_time_minutes"] = metadata["word_count"] / 200
            
            return metadata
            
        except Exception as e:
            print(f"Error extracting metadata: {e}")
            return {
                "word_count": 0,
                "character_count": 0,
                "paragraph_count": 0,
                "sentence_count": 0,
                "average_sentence_length": 0,
                "reading_time_minutes": 0
            }

    def extract_sections(self, content: str) -> Dict[str, str]:
        """Extract different sections from academic document"""
        try:
            sections = {}
            
            # Common academic section patterns
            section_patterns = {
                "abstract": r"abstract\s*[:.]?\s*(.*?)(?=\n\n|\n[A-Z]|$)",
                "introduction": r"introduction\s*[:.]?\s*(.*?)(?=\n\n|\n[A-Z]|$)",
                "methodology": r"methodology\s*[:.]?\s*(.*?)(?=\n\n|\n[A-Z]|$)",
                "results": r"results\s*[:.]?\s*(.*?)(?=\n\n|\n[A-Z]|$)",
                "discussion": r"discussion\s*[:.]?\s*(.*?)(?=\n\n|\n[A-Z]|$)",
                "conclusion": r"conclusion\s*[:.]?\s*(.*?)(?=\n\n|\n[A-Z]|$)",
                "references": r"references?\s*[:.]?\s*(.*?)(?=\n\n|\n[A-Z]|$)",
                "bibliography": r"bibliography\s*[:.]?\s*(.*?)(?=\n\n|\n[A-Z]|$)"
            }
            
            content_lower = content.lower()
            
            for section_name, pattern in section_patterns.items():
                match = re.search(pattern, content_lower, re.IGNORECASE | re.DOTALL)
                if match:
                    sections[section_name] = match.group(1).strip()
            
            return sections
            
        except Exception as e:
            print(f"Error extracting sections: {e}")
            return {}

    def extract_citations(self, content: str) -> list:
        """Extract citations from document content"""
        try:
            citations = []
            
            # Common citation patterns
            citation_patterns = [
                r'\([A-Za-z]+\s+et\s+al\.\s*,\s*\d{4}\)',  # (Author et al., 2023)
                r'\([A-Za-z]+\s+&\s+[A-Za-z]+\s*,\s*\d{4}\)',  # (Author & Author, 2023)
                r'\([A-Za-z]+\s*,\s*\d{4}\)',  # (Author, 2023)
                r'\[[A-Za-z]+\s+et\s+al\.\s*,\s*\d{4}\]',  # [Author et al., 2023]
                r'\[[A-Za-z]+\s+&\s+[A-Za-z]+\s*,\s*\d{4}\]',  # [Author & Author, 2023]
                r'\[[A-Za-z]+\s*,\s*\d{4}\]',  # [Author, 2023]
            ]
            
            for pattern in citation_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                citations.extend(matches)
            
            return list(set(citations))  # Remove duplicates
            
        except Exception as e:
            print(f"Error extracting citations: {e}")
            return []

    def calculate_readability_score(self, content: str) -> Dict[str, float]:
        """Calculate readability scores"""
        try:
            sentences = re.split(r'[.!?]+', content)
            sentences = [s.strip() for s in sentences if s.strip()]
            
            words = content.split()
            syllables = self._count_syllables(content)
            
            if not sentences or not words:
                return {"flesch_reading_ease": 0, "flesch_kincaid_grade": 0}
            
            # Flesch Reading Ease
            flesch_ease = 206.835 - (1.015 * (len(words) / len(sentences))) - (84.6 * (syllables / len(words)))
            flesch_ease = max(0, min(100, flesch_ease))
            
            # Flesch-Kincaid Grade Level
            flesch_grade = (0.39 * (len(words) / len(sentences))) + (11.8 * (syllables / len(words))) - 15.59
            flesch_grade = max(0, flesch_grade)
            
            return {
                "flesch_reading_ease": round(flesch_ease, 2),
                "flesch_kincaid_grade": round(flesch_grade, 2)
            }
            
        except Exception as e:
            print(f"Error calculating readability: {e}")
            return {"flesch_reading_ease": 0, "flesch_kincaid_grade": 0}

    def _count_syllables(self, text: str) -> int:
        """Count syllables in text (simplified)"""
        text = text.lower()
        count = 0
        vowels = "aeiouy"
        on_vowel = False
        
        for char in text:
            is_vowel = char in vowels
            if is_vowel and not on_vowel:
                count += 1
            on_vowel = is_vowel
        
        return max(count, 1)  # At least 1 syllable per word 